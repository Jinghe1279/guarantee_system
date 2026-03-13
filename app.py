from flask import Flask, request, jsonify, send_from_directory, send_file
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import PolynomialFeatures
from sklearn.linear_model import LinearRegression
from sklearn.metrics import mean_squared_error, r2_score
import pandas as pd
import numpy as np
from flask_cors import CORS
from datetime import datetime, date
from decimal import Decimal
import uuid
import joblib
from sklearn.impute import SimpleImputer
# 可选依赖，导入失败时仅影响部分模型功能
try:
    from transformers import BertModel, BertTokenizer  # noqa: F401
except ImportError:
    BertModel = BertTokenizer = None
try:
    import torch  # noqa: F401
except ImportError:
    torch = None
try:
    import shap  # noqa: F401
except ImportError:
    shap = None
try:
    import matplotlib.pyplot as plt  # noqa: F401
except ImportError:
    plt = None
import os
import re
from docx import Document
try:
    from docxtpl import DocxTemplate
except ImportError:
    DocxTemplate = None
from openpyxl import Workbook
from langchain_openai import ChatOpenAI
from langchain_core.messages import SystemMessage, HumanMessage
import logging
from sklearn.preprocessing import StandardScaler
import json
import io
import re
from dotenv import load_dotenv
import pymysql
from werkzeug.security import generate_password_hash, check_password_hash

# Load environment variables
load_dotenv()
os.environ['OPENAI_API_KEY'] = os.getenv('DeepSeek_API_KEY')  # Set API key
# os.environ["OPENAI_API_BASE"] = 'https://api.deepseek.com'  # Set API base URL

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# matplotlib.use('Agg')  # Use Agg backend

app = Flask(__name__)
CORS(app)

# 简单内存令牌存储，重启后自动失效
auth_tokens = {}

# Database configuration
DB_CONFIG = {
    "host": os.getenv("DB_HOST", "localhost"),
    "port": int(os.getenv("DB_PORT", "3306")),
    "user": os.getenv("DB_USER", "root"),
    "password": os.getenv("DB_PASSWORD", "root"),
    "database": os.getenv("DB_NAME", "airport"),
    "cursorclass": pymysql.cursors.DictCursor,
    "charset": "utf8mb4",
}

# Create user table for authentication
def init_user_table():
    try:
        with pymysql.connect(**DB_CONFIG) as conn:
            with conn.cursor() as cursor:
                cursor.execute(
                    """
                    CREATE TABLE IF NOT EXISTS users (
                        id INT AUTO_INCREMENT PRIMARY KEY,
                        username VARCHAR(100) UNIQUE NOT NULL,
                        password_hash VARCHAR(255) NOT NULL,
                        role VARCHAR(20) NOT NULL DEFAULT 'user',
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                    ) CHARACTER SET utf8mb4;
                    """
                )
                # 保证 password_hash 字段足够长，兼容 PBKDF2 哈希
                cursor.execute(
                    """
                    ALTER TABLE users
                    MODIFY password_hash VARCHAR(512) NOT NULL;
                    """
                )
            conn.commit()
        logger.info("User table is ready.")
    except Exception as e:
        logger.error(f"Failed to initialize users table: {e}")

def ensure_admin_user():
    """Ensure there is an admin user; password from env ADMIN_PASSWORD or default 'admin123'."""
    admin_username = os.getenv("ADMIN_USERNAME", "admin")
    admin_password = os.getenv("ADMIN_PASSWORD", "admin123")
    try:
        with pymysql.connect(**DB_CONFIG) as conn:
            with conn.cursor() as cursor:
                cursor.execute("SELECT id, role FROM users WHERE username = %s", (admin_username,))
                user = cursor.fetchone()
                if not user:
                    cursor.execute(
                        "INSERT INTO users (username, password_hash, role) VALUES (%s, %s, %s)",
                        (admin_username, generate_password_hash(admin_password), "admin"),
                    )
                    conn.commit()
                    logger.info("Default admin created.")
                elif user.get("role") != "admin":
                    cursor.execute(
                        "UPDATE users SET role=%s WHERE id=%s",
                        ("admin", user["id"]),
                    )
                    conn.commit()
    except Exception as e:
        logger.error(f"Failed to ensure admin user: {e}")

def get_db_connection():
    """Return a new MySQL connection using shared config."""
    return pymysql.connect(**DB_CONFIG)

init_user_table()
ensure_admin_user()

# 全局模型变量，延迟加载（第一次使用时加载）
ML_MODELS = {
    'scaler_X': None,
    'scaler_y': None,
    'xgb_model': None,
    'scaler_X_binary': None,
    'xgb_model_binary': None,
    'loaded': False,
    'loading': False,  # 防止并发加载
    'xgb_model_failed': False  # 标记主模型是否加载失败
}

def load_ml_models():
    """延迟加载机器学习模型（第一次使用时加载）"""
    if ML_MODELS['loaded']:
        return True
    
    if ML_MODELS['loading']:
        # 如果正在加载，等待
        import time
        while ML_MODELS['loading']:
            time.sleep(0.1)
        return ML_MODELS['loaded']
    
    ML_MODELS['loading'] = True
    
    try:
        logger.info("=" * 60)
        logger.info("开始加载机器学习模型（延迟加载）...")
        logger.info("=" * 60)
        
        # 设置环境变量，避免多进程问题（在加载模型前设置）
        os.environ['OMP_NUM_THREADS'] = '1'
        os.environ['OPENBLAS_NUM_THREADS'] = '1'
        os.environ['MKL_NUM_THREADS'] = '1'
        os.environ['NUMEXPR_NUM_THREADS'] = '1'
        
        logger.info("加载 scaler_X.pkl...")
        ML_MODELS['scaler_X'] = joblib.load('./scaler-model/scaler_X.pkl')
        logger.info("✓ scaler_X.pkl 加载成功")
        
        logger.info("加载 scaler_y.pkl...")
        ML_MODELS['scaler_y'] = joblib.load('./scaler-model/scaler_y.pkl')
        logger.info("✓ scaler_y.pkl 加载成功")
        
        logger.info("加载 scaler_X-binary.pkl...")
        ML_MODELS['scaler_X_binary'] = joblib.load('./scaler-model/scaler_X-binary.pkl')
        logger.info("✓ scaler_X-binary.pkl 加载成功")
        
        logger.info("加载 best_model-binary.pkl...")
        ML_MODELS['xgb_model_binary'] = joblib.load('./scaler-model/best_model-binary.pkl')
        logger.info("✓ best_model-binary.pkl 加载成功")
        
        # 最后加载可能有问题的主模型，如果失败不影响其他模型
        logger.info("加载 best_model.pkl（可能较慢或导致崩溃）...")
        try:
            # 设置环境变量，避免多进程问题（在函数开始时就设置了）
            # 尝试加载主模型，如果失败就跳过
            ML_MODELS['xgb_model'] = joblib.load('./scaler-model/best_model.pkl')
            logger.info("✓ best_model.pkl 加载成功")
            ML_MODELS['xgb_model_failed'] = False
        except Exception as e:
            logger.error(f"加载 best_model.pkl 失败: {e}", exc_info=True)
            logger.warning("主模型加载失败，将仅使用二分类模型（返回0或1）")
            ML_MODELS['xgb_model'] = None
            ML_MODELS['xgb_model_failed'] = True
        except SystemError as e:
            # 捕获可能的段错误相关错误
            logger.error(f"加载 best_model.pkl 时发生系统错误（可能是段错误）: {e}", exc_info=True)
            logger.warning("主模型加载失败（系统错误），将仅使用二分类模型")
            ML_MODELS['xgb_model'] = None
            ML_MODELS['xgb_model_failed'] = True
        
        ML_MODELS['loaded'] = True
        logger.info("=" * 60)
        logger.info("✓ 模型加载完成！")
        logger.info("=" * 60)
        return True
    except FileNotFoundError as e:
        logger.error(f"模型文件未找到: {e}", exc_info=True)
        ML_MODELS['loading'] = False
        return False
    except Exception as e:
        logger.error(f"加载模型时出错: {e}", exc_info=True)
        import traceback
        logger.error(f"加载模型完整错误堆栈: {traceback.format_exc()}")
        ML_MODELS['loading'] = False
        return False
    finally:
        ML_MODELS['loading'] = False

# 不在启动时加载模型，改为延迟加载（第一次使用时加载）
# load_ml_models()  # 注释掉，改为延迟加载

# Configure static folder path for SHAP files
SHAP_DIR = r'C:\Users\liujiaming\Desktop\shantou\Shantou-funding-guarantee-system\shap_files'
if not os.path.exists(SHAP_DIR):
    os.makedirs(SHAP_DIR)

@app.route('/shap_files/<filename>')
def serve_shap_file(filename):
    return send_from_directory(SHAP_DIR, filename)
def convert_numpy_types(obj):
    if isinstance(obj, np.float32):
        return float(obj)
    elif isinstance(obj, dict):
        return {key: convert_numpy_types(value) for key, value in obj.items()}
    elif isinstance(obj, list):
        return [convert_numpy_types(item) for item in obj]
    return obj

def normalize_bool(value):
    if isinstance(value, bool):
        return value
    if value is None:
        return False
    if isinstance(value, (int, float)):
        return value != 0
    if isinstance(value, str):
        val = value.strip()
        if val in ("是", "否"):
            return val == "是"
        val_lower = val.lower()
        if val_lower in ("1", "true", "yes", "y"):
            return True
        if val_lower in ("0", "false", "no", "n", ""):
            return False
    return bool(value)

def normalize_value(value):
    if value is None:
        return ""
    if isinstance(value, (datetime, date)):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, Decimal):
        return str(value)
    return value

def parse_json_field(value):
    if value is None or value == "":
        return []
    if isinstance(value, (list, dict)):
        return value
    if isinstance(value, (bytes, bytearray)):
        try:
            value = value.decode("utf-8")
        except Exception:
            return []
    try:
        return json.loads(value)
    except Exception:
        return []

def normalize_list(value):
    items = parse_json_field(value)
    if not isinstance(items, list):
        return []
    normalized = []
    for item in items:
        if isinstance(item, dict):
            normalized.append({k: normalize_value(v) for k, v in item.items()})
    return normalized

def parse_number(value):
    if value is None or value == "":
        return 0.0
    if isinstance(value, (int, float, Decimal)):
        return float(value)
    if isinstance(value, str):
        text = value.strip().replace(",", "")
        if not text:
            return 0.0
        match = re.search(r"-?\d+(?:\.\d+)?", text)
        if not match:
            return 0.0
        try:
            return float(match.group())
        except ValueError:
            return 0.0
    return 0.0

def sanitize_filename(name):
    text = str(name or "report")
    return re.sub(r'[\\\\/:*?"<>|]+', "_", text)

def build_doc_context(row):
    business_sites_raw = normalize_list(row.get("business_sites_json"))
    business_sites = []
    for site in business_sites_raw:
        normalized = {k: normalize_value(v) for k, v in site.items()}
        normalized["is_pay"] = normalize_bool(site.get("is_pay"))
        business_sites.append(normalized)
    business_accounts = normalize_list(row.get("business_accounts_json"))
    account_rows = normalize_list(row.get("account_rows_json"))
    daily_avg_balance = normalize_list(row.get("daily_avg_balance_json"))
    guarantees = normalize_list(row.get("guarantees_json"))
    existing_loans = normalize_list(row.get("existing_loans_json"))
    electricity_items = normalize_list(row.get("electricity_items_json"))
    asset_stats = normalize_list(row.get("asset_stats_json"))
    rev_check_items = normalize_list(row.get("rev_check_items_json"))
    cashflow_in = normalize_list(row.get("cashflow_in_json"))
    cashflow_out = normalize_list(row.get("cashflow_out_json"))

    month_keys = [f"m{i}" for i in range(1, 13)]

    def parse_optional_number(value):
        if value is None:
            return None
        if isinstance(value, (int, float, Decimal)):
            return float(value)
        if isinstance(value, str):
            text = value.strip().replace(",", "")
            if not text:
                return None
            match = re.search(r"-?\d+(?:\.\d+)?", text)
            if not match:
                return None
            try:
                return float(match.group())
            except ValueError:
                return None
        return None

    def format_total_number(value):
        if value is None:
            return ""
        if abs(value - round(value)) < 1e-9:
            return str(int(round(value)))
        return f"{value:.2f}".rstrip("0").rstrip(".")

    def compare_year_key(text):
        year_text = str(text or "").strip()
        if not year_text:
            return (2, "")
        try:
            return (0, int(year_text))
        except ValueError:
            return (1, year_text)

    totals_by_year = {}
    for item in account_rows:
        year = str((item or {}).get("year") or "").strip()
        if not year:
            continue
        if year not in totals_by_year:
            totals_by_year[year] = {
                "year": year,
                "monthly_totals": {key: 0.0 for key in month_keys},
                "month_has_value": {key: False for key in month_keys},
            }
        year_total = totals_by_year[year]
        for key in month_keys:
            value = parse_optional_number((item or {}).get(key))
            if value is None:
                continue
            year_total["month_has_value"][key] = True
            year_total["monthly_totals"][key] += value

    account_yearly_totals = []
    for year in sorted(totals_by_year.keys(), key=compare_year_key):
        total_item = totals_by_year[year]
        row_total = {"year": year, "avg": ""}
        has_any_month = False
        has_all_months = True
        sum_all = 0.0
        for key in month_keys:
            if not total_item["month_has_value"][key]:
                row_total[key] = ""
                has_all_months = False
                continue
            has_any_month = True
            month_sum = total_item["monthly_totals"][key]
            row_total[key] = format_total_number(month_sum)
            sum_all += month_sum
        if not has_any_month:
            continue
        if has_all_months:
            row_total["avg"] = format_total_number(sum_all / len(month_keys))
        account_yearly_totals.append(row_total)

    def normalize_subject_name(value):
        return value.strip() if isinstance(value, str) else ""

    def has_loan_content(loan):
        if not isinstance(loan, dict):
            return False
        fields = ["type", "amount", "balance", "mode", "monthly_payment", "start_date", "end_date", "bank_rate", "purpose"]
        return any(str(loan.get(key) or "").strip() != "" for key in fields)

    existing_loan_subject_map = {}
    existing_loan_subjects = []
    for loan in existing_loans:
        if not has_loan_content(loan):
            continue
        subject_name = normalize_subject_name((loan or {}).get("loan_subject")) or "未命名贷款主体"
        if subject_name not in existing_loan_subject_map:
            group = {"subject_name": subject_name, "loans": []}
            existing_loan_subject_map[subject_name] = group
            existing_loan_subjects.append(group)
        normalized_loan = {k: normalize_value(v) for k, v in (loan or {}).items()}
        normalized_loan["loan_subject"] = subject_name
        existing_loan_subject_map[subject_name]["loans"].append(normalized_loan)

    existing_amount_total = sum(parse_number(item.get("amount")) for item in existing_loans)
    existing_balance_total = sum(parse_number(item.get("balance")) for item in existing_loans)
    existing_payment_total = sum(parse_number(item.get("monthly_payment")) for item in existing_loans)
    has_existing_values = any(
        str(item.get("amount") or item.get("balance") or item.get("monthly_payment") or "").strip() != ""
        for item in existing_loans
    )
    existing_amount_db = row.get("existing_loans_amount_total")
    existing_balance_db = row.get("existing_loans_balance_total")
    existing_payment_db = row.get("existing_loans_monthly_payment_total")
    has_existing_db = any(
        val not in (None, "") for val in (existing_amount_db, existing_balance_db, existing_payment_db)
    )

    sales_list = normalize_list(row.get("is_table_sales_list_json"))
    if not sales_list:
        legacy_pairs = [
            ("is_table_s1_t", "is_table_s1"),
            ("is_table_s2_t", "is_table_s2"),
            ("is_table_s3_t", "is_table_s3"),
        ]
        for name_key, value_key in legacy_pairs:
            name_val = normalize_value(row.get(name_key))
            value_val = normalize_value(row.get(value_key))
            if name_val or value_val:
                sales_list.append({"name": name_val, "value": value_val})

    context = {
        "project": {
            "a_owner": normalize_value(row.get("project_a_owner")),
            "b_owner": normalize_value(row.get("project_b_owner")),
            "market_manager": normalize_value(row.get("project_market_manager")),
            "source": normalize_value(row.get("project_source")),
            "coop_bank": normalize_value(row.get("project_coop_bank")),
            "apply_date": normalize_value(row.get("project_apply_date")),
            "enterpriseid": normalize_value(row.get("project_enterpriseid")),
        },
        "loan": {
            "borrower_name": normalize_value(row.get("loan_borrower_name")),
            "apply_amount": normalize_value(row.get("loan_apply_amount")),
            "apply_term": normalize_value(row.get("loan_apply_term")),
            "purpose_detail": normalize_value(row.get("loan_purpose_detail")),
        },
        "company": {
            "name": normalize_value(row.get("company_name")),
            "registered_capital": normalize_value(row.get("company_registered_capital")),
            "established_date": normalize_value(row.get("company_established_date")),
            "registered_address": normalize_value(row.get("company_registered_address")),
            "main_business": normalize_value(row.get("company_main_business")),
            "employee_count": normalize_value(row.get("company_employee_count")),
            "is_salary": normalize_bool(row.get("company_is_salary")),
            "shareholder_info": normalize_value(row.get("company_shareholder_info")),
        },
        "controller": {
            "name": normalize_value(row.get("controller_name")),
            "gender": normalize_value(row.get("controller_gender")),
            "native_place": normalize_value(row.get("controller_native_place")),
            "marital_status": normalize_value(row.get("controller_marital_status")),
            "birth_date": normalize_value(row.get("controller_birth_date")),
            "service_years": normalize_value(row.get("controller_service_years")),
            "education": normalize_value(row.get("controller_education")),
            "spouse_name": normalize_value(row.get("controller_spouse_name")),
            "career_experience": normalize_value(row.get("controller_career_experience")),
        },
        "family": {
            "members_info": normalize_value(row.get("family_members_info")),
            "is_hemu": normalize_bool(row.get("family_is_hemu")),
            "annual_expense": normalize_value(row.get("family_annual_expense")),
        },
        "social": {
            "relationship_info": normalize_value(row.get("social_relationship_info")),
        },
        "residence": {
            "type": normalize_value(row.get("residence_type")),
            "years": normalize_value(row.get("residence_years")),
            "address": normalize_value(row.get("residence_address")),
        },
        "business_sites": business_sites,
        "business": {
            "type": normalize_value(row.get("business_type")),
            "month_pay": normalize_value(row.get("business_month_pay")),
            "is_pay": normalize_bool(row.get("business_is_pay")),
            "model_description": normalize_value(row.get("business_model_description")),
            "is_waimao": normalize_bool(row.get("business_is_waimao")),
            "is_jinshen": normalize_bool(row.get("business_is_jinshen")),
        },
        "business_accounts": business_accounts,
        "account_rows": account_rows,
        "account_yearly_totals": account_yearly_totals,
        "daily_avg_balance": daily_avg_balance,
        "guarantees": guarantees,
        "g": {
            "amount_total": normalize_value(row.get("guarantees_amount_total")),
            "balance_total": normalize_value(row.get("guarantees_balance_total")),
        },
        "existing_loans": existing_loans,
        "existing_loan_subjects": existing_loan_subjects,
        "existing_loans_totals": {
            "amount_total": normalize_value(existing_amount_db) if has_existing_db else (f"{existing_amount_total:.2f}" if has_existing_values else ""),
            "balance_total": normalize_value(existing_balance_db) if has_existing_db else (f"{existing_balance_total:.2f}" if has_existing_values else ""),
            "monthly_payment_total": normalize_value(existing_payment_db) if has_existing_db else (f"{existing_payment_total:.2f}" if has_existing_values else ""),
        },
        "credit": {
            "inquiry_count": normalize_value(row.get("credit_inquiry_count")),
            "adverse_info": normalize_value(row.get("credit_adverse_info")),
            "overdue_count": normalize_value(row.get("credit_overdue_count")),
            "max_overdue_amount": normalize_value(row.get("credit_max_overdue_amount")),
        },
        "litigation": {
            "status": normalize_value(row.get("litigation_status")),
        },
        "electricity": {
            "is_quantity": normalize_bool(row.get("electricity_is_quantity")),
            "is_cost": normalize_bool(row.get("electricity_is_cost")),
            "rows": electricity_items,
            "descript": normalize_value(row.get("electricity_descript")),
        },
        "analysis": {
            "plan": {
                "amount": normalize_value(row.get("analysis_plan_amount")),
                "term": normalize_value(row.get("analysis_plan_term")),
                "repayment_method": normalize_value(row.get("analysis_plan_repayment_method")),
                "fee_rate": normalize_value(row.get("analysis_plan_fee_rate")),
                "corp_guarantee": normalize_value(row.get("analysis_plan_corp_guarantee")),
                "personal_guarantee": normalize_value(row.get("analysis_plan_personal_guarantee")),
                "collateral": normalize_value(row.get("analysis_plan_collateral")),
                "diyapingguzhi": normalize_value(row.get("analysis_plan_diyapingguzhi")),
                "eryayuzhi": normalize_value(row.get("analysis_plan_eryayuzhi")),
                "diyajingzhi": normalize_value(row.get("analysis_plan_diyajingzhi")),
            },
            "financials": {
                "total_assets": normalize_value(row.get("analysis_fin_total_assets")),
                "total_liabilities": normalize_value(row.get("analysis_fin_total_liabilities")),
                "net_assets": normalize_value(row.get("analysis_fin_net_assets")),
                "revenue": normalize_value(row.get("analysis_fin_revenue")),
                "net_income": normalize_value(row.get("analysis_fin_net_income")),
            },
            "indicators": {
                "asset_debt_ratio": normalize_value(row.get("analysis_ind_asset_debt_ratio")),
                "sales_debt_ratio": normalize_value(row.get("analysis_ind_sales_debt_ratio")),
                "meets_3x_income": normalize_value(row.get("analysis_ind_meets_3x_income")),
                "receivable_days": normalize_value(row.get("analysis_ind_receivable_days")),
                "avg_balance": normalize_value(row.get("analysis_ind_avg_balance")),
                "repayment_ratio": normalize_value(row.get("analysis_ind_repayment_ratio")),
                "is_superior_loan": normalize_bool(row.get("analysis_ind_is_superior_loan")),
                "is_growth_phase": normalize_bool(row.get("analysis_ind_is_growth_phase")),
                "is_added_guarantor": normalize_bool(row.get("analysis_ind_is_added_guarantor")),
            },
            "soft_info": normalize_value(row.get("analysis_soft_info")),
            "summary": normalize_value(row.get("analysis_summary")),
            "limit": {
                "calculation": normalize_value(row.get("analysis_limit_calculation")),
                "apply_amount": normalize_value(row.get("analysis_limit_apply_amount")),
                "increase_factors": normalize_value(row.get("analysis_limit_increase_factors")),
            },
            "profit_destination": normalize_value(row.get("analysis_profit_destination")),
        },
        "bs": {
            "date": normalize_value(row.get("bs_date")),
            "cash": normalize_value(row.get("bs_cash")),
            "ar": normalize_value(row.get("bs_ar")),
            "prepayments": normalize_value(row.get("bs_prepayments")),
            "other_ar": normalize_value(row.get("bs_other_ar")),
            "inventory": normalize_value(row.get("bs_inventory")),
            "fixed_assets": normalize_value(row.get("bs_fixed_assets")),
            "total_assets": normalize_value(row.get("bs_total_assets")),
            "loans": normalize_value(row.get("bs_loans")),
            "ap": normalize_value(row.get("bs_ap")),
            "advances": normalize_value(row.get("bs_advances")),
            "other_ap": normalize_value(row.get("bs_other_ap")),
            "capital": normalize_value(row.get("bs_capital")),
            "retained_earnings": normalize_value(row.get("bs_retained_earnings")),
            "total_liabilities_equity": normalize_value(row.get("bs_total_liabilities_equity")),
        },
        "asset_stats": asset_stats,
        "asset_totals": {
            "buy_price": normalize_value(row.get("asset_totals_buy_price")),
            "current_value": normalize_value(row.get("asset_totals_current_value")),
            "depreciation": normalize_value(row.get("asset_totals_depreciation")),
        },
        "is_table": {
            "year": normalize_value(row.get("is_table_year")),
            "sales_list": sales_list,
            "s1_t": normalize_value(row.get("is_table_s1_t")),
            "s2_t": normalize_value(row.get("is_table_s2_t")),
            "s3_t": normalize_value(row.get("is_table_s3_t")),
            "s1": normalize_value(row.get("is_table_s1")),
            "s2": normalize_value(row.get("is_table_s2")),
            "s3": normalize_value(row.get("is_table_s3")),
            "s_total": normalize_value(row.get("is_table_s_total")),
            "material_cost": normalize_value(row.get("is_table_material_cost")),
            "gross_profit": normalize_value(row.get("is_table_gross_profit")),
            "f_wages": normalize_value(row.get("is_table_f_wages")),
            "f_rent": normalize_value(row.get("is_table_f_rent")),
            "f_utility": normalize_value(row.get("is_table_f_utility")),
            "f_comm": normalize_value(row.get("is_table_f_comm")),
            "f_trans": normalize_value(row.get("is_table_f_trans")),
            "f_loss": normalize_value(row.get("is_table_f_loss")),
            "f_adv": normalize_value(row.get("is_table_f_adv")),
            "f_entertain": normalize_value(row.get("is_table_f_entertain")),
            "f_tax": normalize_value(row.get("is_table_f_tax")),
            "f_other": normalize_value(row.get("is_table_f_other")),
            "f_total": normalize_value(row.get("is_table_f_total")),
            "net_profit": normalize_value(row.get("is_table_net_profit")),
            "o_family_exp": normalize_value(row.get("is_table_o_family_exp")),
            "annual_net_income": normalize_value(row.get("is_table_annual_net_income")),
            "o_biz_loan": normalize_value(row.get("is_table_o_biz_loan")),
            "o_pvt_loan": normalize_value(row.get("is_table_o_pvt_loan")),
            "o_other_exp": normalize_value(row.get("is_table_o_other_exp")),
            "o_family_inc": normalize_value(row.get("is_table_o_family_inc")),
        },
        "rev_check": {
            "check_list": rev_check_items,
            "total_value": normalize_value(row.get("rev_check_total_value")),
            "est_total": normalize_value(row.get("rev_check_est_total")),
            "is_revenue": normalize_value(row.get("rev_check_is_revenue")),
            "diff_rate": normalize_value(row.get("rev_check_diff_rate")),
            "method": normalize_value(row.get("rev_check_method")),
        },
        "inflow_analysis": cashflow_in,
        "outflow_analysis": cashflow_out,
    }

    return context

def render_docx_from_template(context, template_path):
    doc = DocxTemplate(template_path)
    doc.render(context)

    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)

    final_doc = Document(target_stream)
    for table in final_doc.tables:
        for row in reversed(table.rows):
            if not "".join(cell.text.strip() for cell in row.cells):
                tr = row._tr
                tr.getparent().remove(tr)

    output_stream = io.BytesIO()
    final_doc.save(output_stream)
    output_stream.seek(0)
    return output_stream

@app.route('/register', methods=['POST'])
def register_user():
    data = request.get_json() or {}
    username = data.get('username', '').strip()
    password = data.get('password', '').strip()
    department_id = data.get('department_id')
    # 允许空部门，若有值转为 int
    if department_id not in (None, ''):
        try:
            department_id = int(department_id)
        except (TypeError, ValueError):
            return jsonify({"message": "部门参数不合法"}), 400
    else:
        department_id = None

    if not username or not password:
        return jsonify({"message": "用户名和密码不能为空"}), 400
    if len(password) > 128:
        return jsonify({"message": "密码过长，请少于128字符"}), 400

    try:
        with get_db_connection() as conn:
            with conn.cursor() as cursor:
                # 检查表结构，看是否有 department_id 字段
                try:
                    cursor.execute("SHOW COLUMNS FROM users LIKE 'department_id'")
                    has_dept_field = cursor.fetchone() is not None
                except Exception as e:
                    logger.warning(f"Could not check department_id field: {e}")
                    has_dept_field = False
                
                cursor.execute("SELECT id FROM users WHERE username = %s", (username,))
                if cursor.fetchone():
                    return jsonify({"message": "用户名已存在"}), 400

                # 使用 werkzeug PBKDF2-SHA256 生成哈希，避免 bcrypt 72 字节限制
                password_hash = generate_password_hash(password)
                
                # 根据字段是否存在决定插入语句
                if has_dept_field:
                    cursor.execute(
                        "INSERT INTO users (username, password_hash, role, department_id) VALUES (%s, %s, %s, %s)",
                        (username, password_hash, "user", department_id),
                    )
                else:
                    cursor.execute(
                        "INSERT INTO users (username, password_hash, role) VALUES (%s, %s, %s)",
                        (username, password_hash, "user"),
                    )
            conn.commit()
        return jsonify({"message": "注册成功", "role": "user"}), 201
    except pymysql.Error as e:
        error_msg = str(e)
        logger.error(f"Registration database error: {e}", exc_info=True)
        return jsonify({"message": f"注册失败：数据库错误 - {error_msg}"}), 500
    except Exception as e:
        error_msg = str(e)
        logger.error(f"Registration error: {e}", exc_info=True)
        return jsonify({"message": f"注册失败：{error_msg}"}), 500

@app.route('/login', methods=['POST'])
def login_user():
    data = request.get_json() or {}
    username = data.get('username', '').strip()
    password = data.get('password', '').strip()

    if not username or not password:
        return jsonify({"message": "用户名和密码不能为空"}), 400

    try:
        with get_db_connection() as conn:
            with conn.cursor() as cursor:
                # 检查表结构，看是否有 department_id 字段
                try:
                    cursor.execute("SHOW COLUMNS FROM users LIKE 'department_id'")
                    has_dept_field = cursor.fetchone() is not None
                except Exception as e:
                    logger.warning(f"Could not check department_id field: {e}")
                    has_dept_field = False
                
                # 根据字段是否存在决定查询语句
                if has_dept_field:
                    cursor.execute(
                        "SELECT id, password_hash, role, department_id FROM users WHERE username = %s",
                        (username,),
                    )
                else:
                    cursor.execute(
                        "SELECT id, password_hash, role FROM users WHERE username = %s",
                        (username,),
                    )
                user = cursor.fetchone()
        if not user:
            return jsonify({"message": "用户名或密码错误"}), 401
        if not check_password_hash(user["password_hash"], password):
            return jsonify({"message": "用户名或密码错误"}), 401
        # 确保 admin 账号的角色为 admin
        if username == os.getenv("ADMIN_USERNAME", "admin") and user.get("role") != "admin":
            with get_db_connection() as conn:
                with conn.cursor() as cursor:
                    cursor.execute("UPDATE users SET role=%s WHERE id=%s", ("admin", user["id"]))
                conn.commit()
            user["role"] = "admin"

        # 获取部门名称（如果 department_id 字段存在）
        dept_name = None
        if has_dept_field and user.get("department_id"):
            try:
                with get_db_connection() as conn:
                    with conn.cursor() as cursor:
                        cursor.execute("SELECT name FROM departments WHERE id = %s", (user.get("department_id"),))
                        dept = cursor.fetchone()
                        if dept:
                            dept_name = dept.get("name")
            except Exception as e:
                logger.warning(f"获取部门名称失败: {e}")

        # 生成一次性令牌，后端重启即失效
        token = str(uuid.uuid4())
        auth_tokens[token] = {
            "user_id": user["id"],
            "role": user.get("role", "user"),
            "username": username,
            "created_at": datetime.utcnow().isoformat()
        }

        return jsonify({
            "message": "登录成功",
            "role": user.get("role", "user"),
            "department_name": dept_name,
            "token": token
        })
    except pymysql.Error as e:
        error_msg = str(e)
        logger.error(f"Login database error: {e}", exc_info=True)
        return jsonify({"message": f"登录失败：数据库错误 - {error_msg}"}), 500
    except Exception as e:
        error_msg = str(e)
        logger.error(f"Login error: {e}", exc_info=True)
        return jsonify({"message": f"登录失败：{error_msg}"}), 500

@app.route("/auth-check", methods=["GET"])
def auth_check():
    """校验令牌有效性，用于前端路由守卫"""
    auth_header = request.headers.get("Authorization", "")
    if not auth_header.startswith("Bearer "):
        return jsonify({"message": "未授权"}), 401
    token = auth_header.replace("Bearer ", "").strip()
    info = auth_tokens.get(token)
    if not info:
        return jsonify({"message": "令牌无效或已过期"}), 401
    return jsonify({"message": "ok", "role": info.get("role"), "username": info.get("username")})

class Robot:
    def __init__(self):
        # Initialize the large language model
        self.model = ChatOpenAI(model="deepseek-chat",base_url = "https://api.deepseek.com", temperature=0.2, max_tokens=8192)
        self.row = 0  # Process the first row of input data (single JSON input in this case)

    def qiye_data(self, dataset: dict) -> str:
        """
        Generate a detailed description of the enterprise data for credit assessment.
        :param dataset: Dictionary containing enterprise data
        :return: Formatted report string
        """

        def extract_features(dataset):
            repayment_method_text = dataset.get("repayment_method", "无记录")
            total_monthly_repayment = dataset.get("total_monthly_repayment", 0)
            if total_monthly_repayment in (None, "", 0, 0.0, "0", "0.0"):
                monthly_repayment_text = "未提供"
            else:
                monthly_repayment_text = f"{total_monthly_repayment}万元"
            data = {
                # Project basic information
                "项目编号": dataset.get("project_number", "无记录"),
                "企业名称": dataset.get("company_name", "无记录"),
                "项目经理": dataset.get("project_manager", "无记录"),
                "申请金额": f"{dataset.get('application_amount', 0)}万元",
                "申请期限": f"{dataset.get('application_period', 0)}个月",
                "还款方式": repayment_method_text,
                "每月偿还本息": monthly_repayment_text,

                # Controller information
                "实控人性别": f"{dataset.get('controller_gender', 0)}（0=女，1=男）",
                "文化程度": dataset.get("education_level", "无记录"),
                "婚姻状态": f"{dataset.get('marital_status', 0)}（0=未婚，1=已婚）",
                "居住类型": f"{dataset.get('residence_type', 0)}（0=自购，1=租赁）",
                "本地居住": f"{dataset.get('local_residence_years', 0)}年",

                # Operational characteristics
                "主营业务": dataset.get("main_business", "无记录"),
                "所属行业": dataset.get("industry_category", "无记录"),
                "从业年限": f"{dataset.get('industry_experience', 0)}年",
                "外贸类型": f"{dataset.get('is_foreign_trade', 0)}（0=否，1=是）",
                "谨慎行业": f"{dataset.get('is_cautious_industry', 0)}（0=否，1=是）",
                "员工人数": dataset.get("employee_count", 0),
                "经营场所": f"{dataset.get('business_premises_type', 0)}（0=自有，1=租赁）",
                "月租金": f"{dataset.get('monthly_rent', 0)}万元",

                # Financial data
                "月均余额": f"{dataset.get('monthly_balance', 0)}万元",
                "日均余额": f"{dataset.get('daily_balance', 0)}万元",
                "货币资金": f"{dataset.get('cash_at_meeting', 0)}万元",
                "应收账款": f"{dataset.get('receivables_at_meeting', 0)}万元",
                "存货": f"{dataset.get('inventory_at_meeting', 0)}万元",
                "应付账款": f"{dataset.get('payables_at_meeting', 0)}万元",
                "总资产": f"{dataset.get('total_assets', 0)}万元",
                "总负债": f"{dataset.get('total_liabilities', 0)}万元",
                "净资产": f"{dataset.get('net_assets', 0)}万元",
                "年销售": f"{dataset.get('annual_sales', 0)}万元",
                "年净收益": f"{dataset.get('annual_net_profit', 0)}万元",
                "月净收益": f"{dataset.get('monthly_net_profit', 0)}万元",

                # Risk indicators
                "销售负债率": f"{dataset.get('sales_debt_ratio', 0)}%",
                "资产负债率": f"{dataset.get('asset_debt_ratio', 0)}%",
                "原还款额": f"{dataset.get('monthly_repayment', 0)}万元",
                "新增还款额": f"{dataset.get('total_monthly_repayment', 0)}万元",
                "收益还款比": dataset.get("repayment_income_ratio", 0),

                # Guarantee information
                "房产抵押": f"{dataset.get('property_mortgage', 0)}万元" if dataset.get('property_mortgage',
                                                                                        0) else "无",
                "设备抵押": f"{dataset.get('equipment_mortgage', 0)}万元" if dataset.get('equipment_mortgage',
                                                                                         0) else "无",
                "担保人": f"{dataset.get('additional_guarantor', 0)}（0=否，1=是）",
                "企业保证": f"{dataset.get('company_guarantee', 0)}（0=否，1=是）",
                "个人保证": f"{dataset.get('personal_guarantee', 0)}（0=否，1=是）",

                # Family and credit information
                "家庭和谐": f"{dataset.get('family_harmony', 0)}（1=是，0=否）",
                "未成年子女数": dataset.get("minor_children", 0),
                "成年家庭成员数": dataset.get("adult_family_members", 0),
                "家庭工作成员数": dataset.get("working_family_members", 0),
                "信用查询次数": dataset.get("credit_inquiries", 0),
                "逾期次数": dataset.get("overdue_times", 0),
                "最大逾期金额": dataset.get("max_overdue_amount", 0),
                "银行存入": dataset.get("bank_inflow", 0),
                "银行支出": dataset.get("bank_outflow", 0),
                "最高流入月": f"{dataset.get('highest_flow_month', 0)}",
                "最低流入月": f"{dataset.get('lowest_flow_month', 0)}",

                # Business stage and loan products
                "是否成长阶段": f"{dataset.get('is_growth_stage', 0)}（0=否，1=是）",
                "是否使用优贷宝": f"{dataset.get('used_youdaibao', 0)}（0=否，1=是）",

                # Text information
                "学习及工作经历": dataset.get("education_work_experience", "无记录"),
                "家庭成员及情况": dataset.get("family_social_relations", "无记录"),
                "商业模式": dataset.get("business_model", "无记录"),
                "反担保措施": dataset.get("counter_guarantee", "无记录"),
                "主营业务描述": dataset.get("main_business", "无记录"),
                "近三年利润去向": dataset.get("profit_usage", "无记录"),
                "其他软信息描述": dataset.get("other_soft_info", "无记录"),
                "贷款用途描述": dataset.get("loan_purpose", "无记录"),

                "预测结果": dataset.get('predicted', "无预测结果")
            }

            # Handle missing values
            for k, v in data.items():
                if pd.isna(v) or str(v) == 'nan':
                    data[k] = "无记录"

            return data

        def generate_report(data):
            template = f"""
            **企业信贷评估报告**

            一、基础信息
            项目编号：{data['项目编号']}
            企业名称：{data['企业名称']}（项目经理：{data['项目经理']})
            申请金额：{data['申请金额']}，期限：{data['申请期限']}
            还款方式：{data['还款方式']}
            每月偿还本息：{data['每月偿还本息']}

            二、实控人背景
            1. 性别：{data['实控人性别']}
            2. 文化程度：{data['文化程度']}
            3. 婚姻状态：{data['婚姻状态']}
            4. 居住类型：{data['居住类型']}，本地居住 {data['本地居住']} 年

            三、经营状况
            1. 主营业务：{data['主营业务']}（所属行业：{data['所属行业']}）
            2. 从业年限：{data['从业年限']} 年
            3. 外贸类型：{data['外贸类型']}
            4. 谨慎行业：{data['谨慎行业']}
            5. 员工人数：{data['员工人数']} 人（不含家庭成员）
            6. 经营场所：{data['经营场所']}，月租金 {data['月租金']} 万元

            四、财务概况（单位：万元）
            1. 资金状况：
               - 月均余额：{data['月均余额']}
               - 日均余额：{data['日均余额']}
               - 货币资金：{data['货币资金']}
               - 应收账款：{data['应收账款']}
               - 存货：{data['存货']}
               - 应付账款：{data['应付账款']}

            2. 资产状况：
               - 总资产：{data['总资产']}
               - 总负债：{data['总负债']}
               - 净资产：{data['净资产']}
               - 年销售收入：{data['年销售']}
               - 年净收益：{data['年净收益']}
               - 月净收益：{data['月净收益']}

            五、风险指标
            1. 销售负债率：{data['销售负债率']}
            2. 资产负债率：{data['资产负债率']}
            3. 原贷款月还款额：{data['原还款额']}
            4. 新增月还款额：{data['新增还款额']}
            5. 收益还款比：{data['收益还款比']}

            六、担保措施
            1. 房产抵押：{data['房产抵押']}
            2. 设备抵押：{data['设备抵押']}
            3. 企业保证：{data['企业保证']}
            4. 个人保证：{data['个人保证']}
            5. 是否新增担保人：{data['担保人']}

            七、家庭和信用信息
            1. 家庭和谐：{data['家庭和谐']}
            2. 未成年子女数：{data['未成年子女数']}
            3. 成年家庭成员数：{data['成年家庭成员数']}
            4. 家庭工作成员数：{data['家庭工作成员数']}
            5. 信用查询次数：{data['信用查询次数']}
            6. 逾期次数：{data['逾期次数']}
            7. 最大逾期金额：{data['最大逾期金额']}
            8. 银行存入：{data['银行存入']}
            9. 银行支出：{data['银行支出']}
            10. 最高流入月份：{data['最高流入月']}
            11. 最低流入月份：{data['最低流入月']}

            八、企业发展阶段与贷款产品
            1. 是否成长阶段：{data['是否成长阶段']}
            2. 是否使用优贷宝：{data['是否使用优贷宝']}

            九、文本信息
            1. 学习及工作经历：{data['学习及工作经历']}
            2. 家庭成员及情况：{data['家庭成员及情况']}
            3. 商业模式：{data['商业模式']}
            4. 反担保措施：{data['反担保措施']}
            5. 主营业务描述：{data['主营业务描述']}
            6. 近三年利润去向：{data['近三年利润去向']}
            7. 其他软信息描述：{data['其他软信息描述']}
            8. 贷款用途描述：{data['贷款用途描述']}

            十、其他信息
            1. 预测结果：{data['预测结果']}
            """
            # Clean empty lines while preserving indentation
            return '\n'.join([line for line in template.split('\n') if line.strip()])

        features = extract_features(dataset)
        report = generate_report(features)
        features1 = pd.DataFrame([features])

        return report, features1

    def jingyan_tool(self, file_path: str) -> str:
        """
        Read historical experience from a Word document for reference in credit assessment.
        :param file_path: Path to the Word document
        :return: Text content of the document
        """
        try:
            doc = Document(file_path)
            text_content = "\n".join([para.text for para in doc.paragraphs])
            return text_content
        except Exception as e:
            return f"读取文档时出错: {e}"

    def pinggu_qiye(self, dataset: dict) -> str:
        """
        Predict enterprise credit assessment using machine learning models.
        :param dataset: Dictionary containing enterprise data
        :return: Predicted assessment result
        """
        try:
            # Convert dictionary to DataFrame
            logger.info("开始将字典转换为DataFrame...")
            dataset_df = pd.DataFrame([dataset])
            logger.info(f"DataFrame创建成功，形状: {dataset_df.shape}")
            logger.info(f"Dataset columns: {list(dataset_df.columns)}")

            # Drop text columns not used in numerical prediction
            logger.info("开始删除文本列...")
            columns_to_drop = [
                '主营业务（填写文字信息）', '近三年利润去向描述', '其他软信息描述',
                '学习及工作经历', '家庭成员及情况', '商业模式', '反担保措施',
                '贷款用途描述', '项目结论'
            ]
            dataset_df.drop(columns=[col for col in columns_to_drop if col in dataset_df.columns], inplace=True)
            logger.info(f"删除文本列后，形状: {dataset_df.shape}")

            # Fill missing values with 0
            logger.info("开始填充缺失值...")
            dataset_df = dataset_df.fillna(0)
            logger.info("缺失值填充完成")

            # Required features for the model
            required_features = [
                'application_amount', 'net_assets', 'total_assets', 'core_assets',
                'inventory_at_meeting', 'industry_experience',
                'local_residence_years', 'receivables_at_meeting', 'annual_sales',
                'electricity_consumption', 'total_liabilities', 'repayment_method'
            ]
            
            # Check if all required features exist
            missing_features = [f for f in required_features if f not in dataset_df.columns]
            if missing_features:
                logger.error(f"缺少必需的字段: {missing_features}")
                # Add missing features with default value 0
                for feature in missing_features:
                    dataset_df[feature] = 0
                logger.warning(f"已为缺失字段设置默认值0: {missing_features}")

            # 使用全局加载的模型（避免每次请求都重新加载）
            if not ML_MODELS['loaded']:
                logger.error("模型未加载，尝试重新加载...")
                if not load_ml_models():
                    return "模型加载失败，请检查模型文件"
            
            scaler_X = ML_MODELS['scaler_X']
            scaler_y = ML_MODELS['scaler_y']
            xgb_model = ML_MODELS['xgb_model']
            scaler_X_binary = ML_MODELS['scaler_X_binary']
            xgb_model_binary = ML_MODELS['xgb_model_binary']
            logger.info("使用已加载的模型进行预测")

            # Feature extraction using English column names
            try:
                logger.info(f"开始提取特征，需要的字段: {required_features}")
                X_new = dataset_df[required_features].copy()
                logger.info(f"成功提取特征，形状: {X_new.shape}")
                logger.info(f"提取的特征值: {X_new.iloc[0].to_dict()}")
            except KeyError as e:
                logger.error(f"特征提取失败，缺少字段: {e}", exc_info=True)
                # 尝试使用所有可用的特征
                available_features = [f for f in required_features if f in dataset_df.columns]
                logger.warning(f"使用可用特征: {available_features}")
                if len(available_features) < len(required_features):
                    # 为缺失的特征添加默认值
                    for feature in required_features:
                        if feature not in dataset_df.columns:
                            dataset_df[feature] = 0
                            logger.warning(f"为缺失字段 {feature} 设置默认值0")
                X_new = dataset_df[required_features].copy()
                logger.info(f"重新提取特征后，形状: {X_new.shape}")

            # Now, translate the column names back to Chinese
            logger.info("开始重命名列为中文")
            X_new.columns = [
                '申请金额', '净资产（万元）', '总资产（万元）', '核心资产（固定资产表+货币资金）（万元）',
                '上会时点存货（万元）', '借款人从业限（包括和当前从事行业相关的学习、打工期间）',
                '本地居住时间（年）', '上会时点应收账款（万元）', '年销售收入（万元）',
                '用电量', '总负债（万元）', '申请还款方案还款方式（月还本息之和。单位：万元）'
            ]
            logger.info(f"列重命名完成，当前列名: {list(X_new.columns)}")

            # Scale the features for the binary model
            try:
                logger.info("开始对特征进行标准化（二分类模型）")
                X_new_scaled_bi = scaler_X_binary.transform(X_new)
                logger.info(f"标准化完成，形状: {X_new_scaled_bi.shape}")
                logger.info("开始二分类模型预测")
                pre_bi = xgb_model_binary.predict(X_new_scaled_bi)[0]
                logger.info(f"二分类模型预测结果: {pre_bi}, 类型: {type(pre_bi)}")

                if pre_bi == 1:
                    logger.info("二分类模型预测为通过，开始回归模型预测")
                    # Scale the features for the regression model
                    try:
                        if xgb_model is None or ML_MODELS.get('xgb_model_failed', False):
                            logger.warning("主模型未加载或加载失败，返回二分类结果（通过）")
                            # 返回一个合理的默认值，而不是错误
                            return 1.0  # 返回通过标志
                        
                        X_new_scaled = scaler_X.transform(X_new)
                        logger.info(f"回归模型标准化完成，形状: {X_new_scaled.shape}")
                        predicted_tvbn = xgb_model.predict(X_new_scaled)
                        logger.info(f"回归模型原始预测值: {predicted_tvbn}")
                        predicted_tvbn = scaler_y.inverse_transform(predicted_tvbn.reshape(-1, 1))[0][0]
                        logger.info(f"回归模型最终预测结果: {predicted_tvbn}, 类型: {type(predicted_tvbn)}")
                        return float(predicted_tvbn)
                    except Exception as e:
                        logger.error(f"回归模型预测时出错: {e}", exc_info=True)
                        logger.warning("回归模型预测失败，返回二分类结果（通过）")
                        return 1.0  # 返回通过标志，而不是错误
                else:
                    logger.info(f"二分类模型预测为不予过会: {pre_bi}")
                    return float(pre_bi) if isinstance(pre_bi, (int, float)) else 0
            except Exception as e:
                logger.error(f"模型预测时出错: {e}", exc_info=True)
                return f"模型预测失败: {str(e)}"
                
        except Exception as e:
            logger.error(f"pinggu_qiye函数执行出错: {e}", exc_info=True)
            return f"预测过程出错: {str(e)}"


def map_database_fields_to_model_fields(db_record: dict) -> dict:
    """
    将数据库字段名映射为模型期望的字段名
    :param db_record: 数据库原始记录（下划线字段名）
    :return: 映射后的字典（英文字段名）
    """
    # 字段映射表：数据库字段名 -> 模型字段名
    field_mapping = {
        # 项目信息
        'project_enterpriseid': 'project_number',
        'project_market_manager': 'project_manager',
        'project_a_owner': 'project_manager',  # 如果没有market_manager，使用a_owner
        
        # 贷款信息
        'loan_apply_amount': 'application_amount',
        'loan_apply_term': 'application_period',
        'loan_purpose_detail': 'loan_purpose',
        'loan_borrower_name': 'borrower_name',
        
        # 公司信息
        'company_name': 'company_name',
        'company_main_business': 'main_business',
        'company_employee_count': 'employee_count',
        
        # 实控人信息
        'controller_name': 'controller_name',
        'controller_gender': 'controller_gender',
        'controller_education': 'education_level',
        'controller_marital_status': 'marital_status',
        'controller_service_years': 'industry_experience',
        'controller_career_experience': 'education_work_experience',
        'controller_birth_date': 'birth_date',
        'controller_native_place': 'native_place',
        'controller_spouse_name': 'spouse_name',
        
        # 居住信息
        'residence_type': 'residence_type',
        'residence_years': 'local_residence_years',
        'residence_address': 'residence_address',
        
        # 经营信息
        'business_type': 'business_type',
        'business_model_description': 'business_model',
        'business_is_waimao': 'is_foreign_trade',
        'business_is_jinshen': 'is_cautious_industry',
        'business_month_pay': 'monthly_rent',
        'business_is_pay': 'business_premises_type',  # 0=自有，1=租赁
        
        # 财务信息
        'analysis_fin_total_assets': 'total_assets',
        'analysis_fin_total_liabilities': 'total_liabilities',
        'analysis_fin_net_assets': 'net_assets',
        'analysis_fin_revenue': 'annual_sales',
        'analysis_fin_net_income': 'annual_net_profit',
        
        # 资产负债表字段
        'bs_cash': 'cash_at_meeting',
        'bs_ar': 'receivables_at_meeting',
        'bs_inventory': 'inventory_at_meeting',
        'bs_ap': 'payables_at_meeting',
        
        # 指标
        'analysis_ind_asset_debt_ratio': 'asset_debt_ratio',
        'analysis_ind_sales_debt_ratio': 'sales_debt_ratio',
        'analysis_ind_avg_balance': 'monthly_balance',
        'analysis_ind_receivable_days': 'receivable_days',
        'analysis_ind_is_growth_phase': 'is_growth_stage',
        'analysis_ind_is_added_guarantor': 'additional_guarantor',
        
        # 还款方案
        'analysis_plan_repayment_method': 'repayment_method',
        
        # 其他
        'analysis_profit_destination': 'profit_usage',
        'analysis_soft_info': 'other_soft_info',
        'family_members_info': 'family_social_relations',
        'credit_inquiry_count': 'credit_inquiries',
        'credit_overdue_count': 'overdue_times',
        'credit_max_overdue_amount': 'max_overdue_amount',
        
        # 用电量 - 需要从electricity_items_json计算
        # 'electricity_consumption': 需要计算
        
        # 行业类别 - 可能需要从其他字段推断
        # 'industry_category': 需要推断
    }
    
    # 创建映射后的字典
    mapped_data = {}
    
    # 直接映射的字段
    for db_field, model_field in field_mapping.items():
        if db_field in db_record and db_record[db_field] is not None:
            mapped_data[model_field] = db_record[db_field]
    
    # 特殊处理：project_manager
    if 'project_manager' not in mapped_data:
        if db_record.get('project_market_manager'):
            mapped_data['project_manager'] = db_record['project_market_manager']
        elif db_record.get('project_a_owner'):
            mapped_data['project_manager'] = db_record['project_a_owner']
    
    # 特殊处理：business_premises_type (0=自有，1=租赁)
    if 'business_is_pay' in db_record:
        is_pay = db_record['business_is_pay']
        if isinstance(is_pay, str):
            mapped_data['business_premises_type'] = 1 if is_pay in ('是', '1', 'true', 'True') else 0
        else:
            mapped_data['business_premises_type'] = 1 if is_pay else 0
    
    # 特殊处理：is_foreign_trade
    if 'business_is_waimao' in db_record:
        is_waimao = db_record['business_is_waimao']
        if isinstance(is_waimao, str):
            mapped_data['is_foreign_trade'] = 1 if is_waimao in ('是', '1', 'true', 'True') else 0
        else:
            mapped_data['is_foreign_trade'] = 1 if is_waimao else 0
    
    # 特殊处理：is_cautious_industry
    if 'business_is_jinshen' in db_record:
        is_jinshen = db_record['business_is_jinshen']
        if isinstance(is_jinshen, str):
            mapped_data['is_cautious_industry'] = 1 if is_jinshen in ('是', '1', 'true', 'True') else 0
        else:
            mapped_data['is_cautious_industry'] = 1 if is_jinshen else 0
    
    # 特殊处理：controller_gender (转换为0/1)
    if 'controller_gender' in mapped_data:
        gender = mapped_data['controller_gender']
        if isinstance(gender, str):
            mapped_data['controller_gender'] = 1 if gender in ('男', '1', 'M', 'male') else 0
        elif gender is None:
            mapped_data['controller_gender'] = 0
    
    # 特殊处理：marital_status (转换为0/1)
    if 'marital_status' in mapped_data:
        marital = mapped_data['marital_status']
        if isinstance(marital, str):
            mapped_data['marital_status'] = 1 if marital in ('已婚', '1', 'married') else 0
        elif marital is None:
            mapped_data['marital_status'] = 0
    
    # 特殊处理：residence_type (转换为0/1，0=自购，1=租赁)
    if 'residence_type' in mapped_data:
        residence = mapped_data['residence_type']
        if isinstance(residence, str):
            mapped_data['residence_type'] = 1 if residence in ('租赁', '1', 'rent') else 0
        elif residence is None:
            mapped_data['residence_type'] = 0
    
    # 计算用电量（从electricity_items_json）
    if 'electricity_items_json' in db_record and db_record['electricity_items_json']:
        try:
            import json
            if isinstance(db_record['electricity_items_json'], str):
                electricity_items = json.loads(db_record['electricity_items_json'])
            else:
                electricity_items = db_record['electricity_items_json']
            
            if isinstance(electricity_items, list) and len(electricity_items) > 0:
                # 计算最近一年的总用电量
                total_consumption = 0
                for item in electricity_items:
                    if isinstance(item, dict) and 'total' in item:
                        try:
                            total_consumption += float(item['total'] or 0)
                        except (ValueError, TypeError):
                            pass
                mapped_data['electricity_consumption'] = total_consumption
        except Exception as e:
            logger.warning(f"Failed to calculate electricity_consumption: {e}")
            mapped_data['electricity_consumption'] = 0
    
    # 计算核心资产（固定资产+货币资金）
    if 'bs_fixed_assets' in db_record and 'bs_cash' in db_record:
        try:
            fixed_assets = float(db_record.get('bs_fixed_assets') or 0)
            cash = float(db_record.get('bs_cash') or 0)
            mapped_data['core_assets'] = fixed_assets + cash
        except (ValueError, TypeError):
            mapped_data['core_assets'] = 0
    
    # 计算月还款额（从existing_loans_json）
    if 'existing_loans_json' in db_record and db_record['existing_loans_json']:
        try:
            import json
            if isinstance(db_record['existing_loans_json'], str):
                existing_loans = json.loads(db_record['existing_loans_json'])
            else:
                existing_loans = db_record['existing_loans_json']
            
            if isinstance(existing_loans, list):
                total_monthly_payment = 0
                for loan in existing_loans:
                    if isinstance(loan, dict) and isinstance(loan.get('loans'), list):
                        for child_loan in loan.get('loans', []):
                            if isinstance(child_loan, dict) and 'monthly_payment' in child_loan:
                                try:
                                    total_monthly_payment += float(child_loan['monthly_payment'] or 0)
                                except (ValueError, TypeError):
                                    pass
                        continue
                    if isinstance(loan, dict) and 'monthly_payment' in loan:
                        try:
                            total_monthly_payment += float(loan['monthly_payment'] or 0)
                        except (ValueError, TypeError):
                            pass
                mapped_data['monthly_repayment'] = total_monthly_payment
        except Exception as e:
            logger.warning(f"Failed to calculate monthly_repayment: {e}")
    
    # 计算新增还款额（申请还款方案）
    if 'analysis_plan_repayment_method' in db_record:
        try:
            repayment_value = db_record.get('analysis_plan_repayment_method')
            if isinstance(repayment_value, (int, float, Decimal)):
                mapped_data['total_monthly_repayment'] = float(repayment_value)
            elif isinstance(repayment_value, str) and repayment_value.strip():
                mapped_data['total_monthly_repayment'] = float(repayment_value)
        except (ValueError, TypeError):
            pass
    
    # 计算收益还款比
    if 'monthly_net_profit' in mapped_data or 'is_table_annual_net_income' in db_record:
        try:
            monthly_profit = mapped_data.get('monthly_net_profit', 0)
            if not monthly_profit and 'is_table_annual_net_income' in db_record:
                annual_income = float(db_record.get('is_table_annual_net_income') or 0)
                monthly_profit = annual_income / 12
            
            total_repayment = mapped_data.get('total_monthly_repayment', 0)
            if total_repayment > 0:
                mapped_data['repayment_income_ratio'] = monthly_profit / total_repayment
            else:
                mapped_data['repayment_income_ratio'] = 0
        except (ValueError, TypeError, ZeroDivisionError):
            mapped_data['repayment_income_ratio'] = 0
    
    # 计算月净收益
    if 'is_table_annual_net_income' in db_record:
        try:
            annual_income = float(db_record.get('is_table_annual_net_income') or 0)
            mapped_data['monthly_net_profit'] = annual_income / 12
        except (ValueError, TypeError):
            mapped_data['monthly_net_profit'] = 0
    
    # 日均余额（从daily_avg_balance_json计算）
    if 'daily_avg_balance_json' in db_record and db_record['daily_avg_balance_json']:
        try:
            import json
            if isinstance(db_record['daily_avg_balance_json'], str):
                daily_balance_data = json.loads(db_record['daily_avg_balance_json'])
            else:
                daily_balance_data = db_record['daily_avg_balance_json']
            
            if isinstance(daily_balance_data, list) and len(daily_balance_data) > 0:
                # 取最近一年的年均余额
                latest_year = daily_balance_data[-1]
                if isinstance(latest_year, dict) and 'annual_avg' in latest_year:
                    mapped_data['daily_balance'] = float(latest_year.get('annual_avg') or 0)
        except Exception as e:
            logger.warning(f"Failed to calculate daily_balance: {e}")
    
    # 行业类别 - 从main_business推断或使用默认值
    if 'industry_category' not in mapped_data:
        # 可以尝试从main_business推断，这里先设为空
        mapped_data['industry_category'] = db_record.get('company_main_business', '')[:50] if db_record.get('company_main_business') else ''
    
    # 确保所有模型必需的字段都有默认值
    required_fields = {
        'application_amount': 0,
        'net_assets': 0,
        'total_assets': 0,
        'core_assets': 0,
        'inventory_at_meeting': 0,
        'industry_experience': 0,
        'local_residence_years': 0,
        'receivables_at_meeting': 0,
        'annual_sales': 0,
        'electricity_consumption': 0,
        'total_liabilities': 0,
        'repayment_method': 0,
        'controller_gender': 0,
        'marital_status': 0,
        'residence_type': 0,
        'is_foreign_trade': 0,
        'is_cautious_industry': 0,
        'business_premises_type': 0,
        'monthly_balance': 0,
        'daily_balance': 0,
        'monthly_rent': 0,
        'cash_at_meeting': 0,
        'payables_at_meeting': 0,
        'annual_net_profit': 0,
        'monthly_net_profit': 0,
        'asset_debt_ratio': 0,
        'sales_debt_ratio': 0,
        'monthly_repayment': 0,
        'total_monthly_repayment': 0,
        'repayment_income_ratio': 0,
    }
    
    # 为缺失的必需字段设置默认值
    for field, default_value in required_fields.items():
        if field not in mapped_data or mapped_data[field] is None:
            mapped_data[field] = default_value
        # 确保数值类型正确
        try:
            if isinstance(mapped_data[field], str):
                # 尝试转换为数字
                mapped_data[field] = float(mapped_data[field]) if mapped_data[field].strip() else default_value
            else:
                mapped_data[field] = float(mapped_data[field]) if mapped_data[field] is not None else default_value
        except (ValueError, TypeError):
            mapped_data[field] = default_value
    
    return mapped_data

@app.route('/demo', methods=['POST'])
def data_type():
    data = request.get_json()
    print(data)#post方法接收到的参数
    logger.info(f"Received data: {data}")

    # 将数据库字段名映射为模型期望的字段名
    mapped_data = map_database_fields_to_model_fields(data)
    logger.info(f"Mapped data: {mapped_data}")

    robot = Robot()
    history_file = "./RAG/贷款评审要素_总结版.docx"

    # Load historical experience
    history_data = robot.jingyan_tool(history_file)

    # Generate enterprise data report (使用映射后的数据)
    try:
        qi_ye_data,dataset = robot.qiye_data(mapped_data)
        print(qi_ye_data)#特征名
        print(dataset)#带具体特征的prompt
    except Exception as e:
        logger.error(f"生成企业数据报告时出错: {e}", exc_info=True)
        return jsonify({"message": f"生成报告失败: {str(e)}", "模型预测金额": None, "大模型判断": None}), 500
    
    # Predict using machine learning models (使用映射后的数据)
    try:
        logger.info("开始调用模型预测函数")
        model_result = robot.pinggu_qiye(mapped_data)
        logger.info(f"模型预测函数返回: {model_result}, 类型: {type(model_result)}")
        
        # 检查预测结果是否为错误字符串
        if isinstance(model_result, str):
            if "错误" in model_result or "失败" in model_result or "未找到" in model_result:
                logger.error(f"模型预测返回错误字符串: {model_result}")
                return jsonify({
                    "message": "模型预测失败",
                    "模型预测金额": None,
                    "大模型判断": None,
                    "错误详情": model_result
                }), 500
            else:
                # 尝试将字符串转换为数值
                try:
                    model_result = float(model_result)
                    logger.info(f"成功将字符串转换为数值: {model_result}")
                except (ValueError, TypeError):
                    logger.error(f"模型预测结果无法转换为数值: {model_result}")
                    return jsonify({
                        "message": "模型预测结果格式错误",
                        "模型预测金额": None,
                        "大模型判断": None,
                        "错误详情": f"预测结果: {model_result}"
                    }), 500
        
        # 确保 model_result 是数值类型
        try:
            if not isinstance(model_result, (int, float)):
                model_result = float(model_result)
            logger.info(f"最终模型预测结果: {model_result}, 类型: {type(model_result)}")
        except (ValueError, TypeError) as e:
            logger.error(f"模型预测结果无法转换为数值: {model_result}, 错误: {e}")
            return jsonify({
                "message": "模型预测结果格式错误",
                "模型预测金额": None,
                "大模型判断": None,
                "错误详情": f"预测结果: {model_result}, 类型: {type(model_result)}"
            }), 500
            
    except Exception as e:
        logger.error(f"模型预测过程出错: {e}", exc_info=True)
        import traceback
        logger.error(f"完整错误堆栈: {traceback.format_exc()}")
        return jsonify({
            "message": "模型预测失败",
            "模型预测金额": None,
            "大模型判断": None,
            "错误详情": str(e)
        }), 500

    # Prepare prompt for LLM
    prompt = f"""                
                **系统角色：**
                ```
                您是一位经验丰富的中小微企业融资担保评估专家，擅长分析企业财务数据和描述信息，识别影响融资担保评估准确性的关键因素。
                ```
            你的任务是：

            1. 根据用户的问题，调用适当的工具（如RAG工具或API），提供最准确的企业资产预测和查询信息。若用户要进行融资担保评估，则首先要接收模型对该企业的融资担保评估的预测结果，还要根据企业的信息，和你的知识储备，
            以及下面给出的历史评估经验，综合给用户提意见，即是否向该企业贷款或贷款金额是否合适或还需要重点关注该企业的哪些问题。注意对于评估经验文本要仔细学习，按照每条的要求对输入的数据进行分析，将分析过程以及最终结果和结合模型给出的最终建议的评估金额全部输出，要求非常的详细。
            注意：机器学习模型给出的本笔贷款可申请金额的参考价值占到80%，要在预测结果的基础上，结合经验和相关知识等，判断是否批准这笔贷款担保申请。
            历史评估经验如下："{history_data}".  
            2. 在回答时：
                - 保持语气温和、友好，并提供有用的建议。
                - 请分步骤说明复杂的查询结果，必要时详细解释背景知识。
                - 确保答案简洁、清晰，但必要时可以详细解释背景知识。

            3. 处理特殊情况
                - 如果数据不可用或无法确定结果，请礼貌地告知用户原因，并建议用户提供更多信息（如具体企业ID）。
                - 如果用户输入的格式不正确，请友好地提醒并提供正确的输入格式。
            4. **应用总结经验进行新数据评估：**
    ·		•步骤：
                分析新企业的描述信息，聚焦与历史错误预测相关的特征。
                依次检索以往数据与经验文档，检查可参照的类似案例和核心特征组合。
                结合特征贡献度排名，若在特征间出现冲突则倾向依赖排名更高的特征做最终判定。
                做出结论：
    •		如果对模型预测完全认可，输出“建议批准本笔贷款担保申请”；
    •		如果无法确定或认为模型的预测结果有误，输出“建议暂缓本笔贷款担保申请”。

        
            **注意事项：**
            1. **逐步推理（思维链）：**
                - 在内部进行详细分析和推理，但最终回答中不展示这些步骤。
            2. **准确性和专业性：**
                - 完全依赖于历史数据中总结的规律和经验，避免主观猜测。
            3. **明确和绝对的结论：**
                - 结论应明确无误，避免使用模棱两可的语言。
            4. **提升模型准确率：**
                - 关注预测模型的86%准确率，因此要有百分百的把握给出结论，不怕思考时间长，要反复检查思考最后再得出结论。
            5. **重点挖掘误差原因：**
                - 深入分析导致预测错误的具体特征或特征组合，提出避免误差的建议。
        
        """
    # 从映射后的数据中获取关键信息，用于生成提示文本
    application_amount = mapped_data.get('application_amount', 0)
    application_period = mapped_data.get('application_period', 0)
    repayment_method = mapped_data.get('repayment_method') or "未提供"
    monthly_repayment = mapped_data.get('total_monthly_repayment', 0)

    def format_llm_number(value, digits=2):
        try:
            number = float(value)
            return f"{number:.{digits}f}".rstrip("0").rstrip(".")
        except (TypeError, ValueError):
            return str(value)

    application_amount_text = format_llm_number(application_amount)
    application_period_text = str(application_period)
    model_result_text = format_llm_number(model_result)
    if monthly_repayment in (None, "", 0, 0.0, "0", "0.0"):
        monthly_repayment_text = "未提供"
    else:
        monthly_repayment_text = f"{format_llm_number(monthly_repayment)}万元"

    def clean_llm_judgment(text: str, application_period_value) -> str:
        cleaned = (text or "").replace("\r\n", "\n").replace("\r", "\n")
        replacements = [
            ("请按照给出的说明进行操作。一步一步地仔细分析所提供的数据（内部分析，不要暴露你的思维链），然后按照下列要求的格式提出你的最终结论。", ""),
            ("**最终输出格式需要包含以下5点内容：**", ""),
            ("（样本中申请金额的值）", ""),
            ("（样本中申请期限的值）", ""),
            ("（样本中申请还款方案还款方式（月还本息之和。单位：万元）的值）", ""),
            ("（这部分根据输入的样本中内容输出）", ""),
            ("（这部分根据给出的模型预测结果的值输出）", ""),
            ("（这部分根据判断的结果输出）", ""),
            ("详细说明判断的依据和关键因素，并且给出详细建议和评审理由。", ""),
        ]
        for old, new in replacements:
            cleaned = cleaned.replace(old, new)

        cleaned = re.sub(r"\[建议批准本笔贷款担保申请/ ?建议暂缓本笔贷款担保申请\]", "", cleaned)
        cleaned = re.sub(r"- 4\..*?(?=\n- 5\.|\Z)", "", cleaned, flags=re.S)
        cleaned = re.sub(r"```+", "", cleaned)
        cleaned = re.sub(rf"期限：\s*{re.escape(str(application_period_value))}年", f"期限：{application_period_value}个月", cleaned)
        cleaned = re.sub(r"[ \t]+\n", "\n", cleaned)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        return cleaned.strip()
    
    # Prepare text for LLM based on model prediction
    if model_result == 0:
        text = f"""
新待判断数据：
预测模型认为该公司不予过会。企业数据如下：
{qi_ye_data}

请只输出最终结论，不要输出任何提示词说明、模板说明、括号备注、方括号占位符、反引号代码块，也不要复述“按以下格式输出”等指令。

输出要求：
1. 使用“个月”，不要使用“年”描述申请期限。
2. “还款方式”与“每月偿还本息”是两个不同字段，不要混淆。
3. 如果无法从输入中确定“每月偿还本息”，明确写“未提供”，不要编造，也不要写 0 万元。
4. 仅按以下结构输出：

1. 该企业申请的评估信息：
申请金额：{application_amount_text}万元
期限：{application_period_text}个月
还款方式：{repayment_method}
每月偿还本息：{monthly_repayment_text}

2. 模型对该企业的评估资产预测结果：
{model_result_text}万元

3. 最终结论：
只能二选一：
建议批准本笔贷款担保申请
或
建议暂缓本笔贷款担保申请

4. 结论摘要：
用自然语言详细说明判断依据、关键因素、风险点与建议。
"""
    else:
        text = f"""
新待判断数据：
预测模型对该公司的预测结果如下：{model_result_text}万元。企业数据如下：
{qi_ye_data}

请只输出最终结论，不要输出任何提示词说明、模板说明、括号备注、方括号占位符、反引号代码块，也不要复述“按以下格式输出”等指令。

输出要求：
1. 使用“个月”，不要使用“年”描述申请期限。
2. “还款方式”与“每月偿还本息”是两个不同字段，不要混淆。
3. 如果无法从输入中确定“每月偿还本息”，明确写“未提供”，不要编造，也不要写 0 万元。
4. 仅按以下结构输出：

1. 该企业申请的评估信息：
申请金额：{application_amount_text}万元
期限：{application_period_text}个月
还款方式：{repayment_method}
每月偿还本息：{monthly_repayment_text}

2. 模型对该企业的评估资产预测结果：
{model_result_text}万元

3. 最终结论：
只能二选一：
建议批准本笔贷款担保申请
或
建议暂缓本笔贷款担保申请

4. 结论摘要：
用自然语言详细说明判断依据、关键因素、风险点与建议。
"""

    # Call the large language model for final decision
    try:
        final_decision = robot.model.invoke([
        SystemMessage(content=prompt),
        HumanMessage(content=f"输入数据: {text}。")
    ])
        final_judgment = clean_llm_judgment(final_decision.content, application_period_text)
        print(final_judgment)
    except Exception as e:
        logger.error(f"调用大模型时出错：{e}")
        final_judgment = "大模型评估失败"

    # Return the response
    response = {
    "模型预测金额": model_result,
    "大模型判断": final_judgment
    }
    response = convert_numpy_types(response)
    return jsonify(response)

@app.route('/download-report', methods=['GET'])
def download_report():
    application_id = (request.args.get('id') or '').strip()
    if not application_id:
        logger.error("下载报告：缺少ID参数")
        return jsonify({"message": "missing id"}), 400
    
    if DocxTemplate is None:
        logger.error("下载报告：docxtpl未安装")
        return jsonify({"message": "docxtpl not installed"}), 500

    template_path = os.path.join(os.path.dirname(__file__), "template", "template.docx")
    if not os.path.exists(template_path):
        logger.error(f"下载报告：模板文件不存在: {template_path}")
        return jsonify({"message": f"template file not found: {template_path}"}), 500

    try:
        logger.info(f"下载报告：开始查询记录，ID: {application_id}")
        with get_db_connection() as conn:
            with conn.cursor() as cursor:
                cursor.execute("SELECT * FROM loan_application WHERE id = %s", (application_id,))
                row = cursor.fetchone()
        
        if not row:
            logger.error(f"下载报告：记录未找到，ID: {application_id}")
            return jsonify({"message": f"record not found for id: {application_id}"}), 404

        logger.info(f"下载报告：记录查询成功，开始构建文档上下文")
        try:
            context = build_doc_context(row)
            logger.info(f"下载报告：文档上下文构建成功")
        except Exception as e:
            logger.error(f"下载报告：构建文档上下文失败: {e}", exc_info=True)
            return jsonify({"message": f"failed to build document context: {str(e)}"}), 500

        logger.info(f"下载报告：开始渲染文档")
        try:
            output_stream = render_docx_from_template(context, template_path)
            logger.info(f"下载报告：文档渲染成功")
        except Exception as e:
            logger.error(f"下载报告：渲染文档失败: {e}", exc_info=True)
            return jsonify({"message": f"failed to render document: {str(e)}"}), 500

        safe_name = sanitize_filename(row.get("project_enterpriseid") or row.get("company_name") or f"report-{application_id}")
        filename = f"{safe_name}-report.docx"
        logger.info(f"下载报告：准备返回文件，文件名: {filename}")
        
        return send_file(
            output_stream,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except pymysql.Error as e:
        logger.error(f"下载报告：数据库错误: {e}", exc_info=True)
        return jsonify({"message": f"database error: {str(e)}"}), 500
    except Exception as e:
        logger.error(f"下载报告：生成报告失败: {e}", exc_info=True)
        import traceback
        logger.error(f"下载报告：完整错误堆栈: {traceback.format_exc()}")
        return jsonify({"message": f"failed to generate report: {str(e)}"}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', debug=True, port=int(os.getenv("FLASK_PORT", "5001")))
